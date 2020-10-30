from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt

def write_table(document, headers, data):
	
	hdr_border = {"sz": 6, "val": "single", "color": "#000000", "space": "0"} 
	
	table = document.add_table(rows=1, cols=len(headers)+1)
	
	#Add Headers
	hdr_cells = table.rows[0].cells
	for index, h in enumerate(headers):
		print(h)
		hdr_cells[index].text = h
	
		set_cell_border(hdr_cells[index], top=hdr_border,  bottom=hdr_border, start=hdr_border, end=hdr_border)
		
		shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b1b5b2"/>'.format(nsdecls('w')))
		hdr_cells[index]._tc.get_or_add_tcPr().append(shading_elm_1)
		
	for row in data:
		new_row = table.add_row()
		row_cells = new_row.cells	
		
		for index, c in enumerate(row):
			run = row_cells[index].paragraphs[0].add_run(str(c))
			set_cell_border(row_cells[index], bottom = hdr_border, start=hdr_border, end=hdr_border)
			
def set_cell_border(cell, **kwargs):
	"""
	Set cell`s border
	Usage:

	set_cell_border(
		cell,
		top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
		bottom={"sz": 12, "color": "#00FF00", "val": "single"},
		start={"sz": 24, "val": "dashed", "shadow": "true"},
		end={"sz": 12, "val": "dashed"},
	)
	"""
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()

	# check for tag existnace, if none found, then create one
	tcBorders = tcPr.first_child_found_in("w:tcBorders")
	if tcBorders is None:
		tcBorders = OxmlElement('w:tcBorders')
		tcPr.append(tcBorders)

	# list over all available tags
	for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
		edge_data = kwargs.get(edge)
		if edge_data:
			tag = 'w:{}'.format(edge)

			# check for tag existnace, if none found, then create one
			element = tcBorders.find(qn(tag))
			if element is None:
				element = OxmlElement(tag)
				tcBorders.append(element)

			# looks like order of attributes is important
			for key in ["sz", "val", "color", "space", "shadow"]:
				if key in edge_data:
					element.set(qn('w:{}'.format(key)), str(edge_data[key]))