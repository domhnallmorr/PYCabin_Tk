from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt

class Service_Bulletin():

	def __init__(self, mainapp, change):
		
		self.change = change
		self.mainapp = mainapp
		
		self.file_path = r'C:\Users\domhn\Documents\Python\Pycabin_Tkinter\V0.08\SB_Test.docx'
		
		self.determine_ac_groups()
		
		self.gen_parts_table()
		
		self.gen_document_references()
		
		self.get_oem_docs()
		
	def determine_ac_groups(self):
		
		#group in format {'MSN', 'A/C Type'}
		
		self.groups = {}
		
		ac = []
		
		for mod in self.change.mods:
			for a in mod[-1].split(','):
				if a not in ac:
					ac.append(a)
		
		self.groups = {1: ac}
		
		
	def gen_parts_table(self):
		
		self.parts_table = []
		for mod in self.change.mods:
			mod = self.mainapp.frames[mod[1]].backend
			
			pt = mod.gen_parts_table()
			
			for p in pt:
				self.parts_table.append(p)

	
		
	def gen_document_references(self):
	
		self.drawings = self.change.get_change_drawings()
		
		for mod in self.change.mods:
			if mod[0] == 'LOPA':
				#get seat CMM
				pass
				
	def get_oem_docs(self):
		print('here')
		self.oem_docs = self.change.get_change_oem_docs()
	
	def gen_steps_required(self):
	
		pass
		


class Service_Bulletin_Writer():

	def __init__(self, sb):
	
		self.sb = sb
		
		self.setup_styles()
		
		self.create_doc()
		self.write_parts_section()
		self.write_references_section()
		
		self.document.save(self.sb.file_path)
		
	def create_doc(self):
	
		self.document = Document(r'C:\Users\domhn\Documents\Python\Pycabin_Tkinter\Resources\archive\SB_Template.docx')

	def setup_styles(self):
	
		self.hdr_border = {"sz": 6, "val": "single", "color": "#000000", "space": "0"} 
		
	def write_parts_section(self):
		
		self.document.add_heading('Provisioning', level=1)	

		self.write_table(['Qty', 'Part Number', 'Description'], self.sb.parts_table)
		
		p = self.document.add_paragraph()
		run = p.add_run()
		run.add_break(WD_BREAK.PAGE)
	
	def write_references_section(self):
	
		self.document.add_heading('References', level=1)
		self.document.add_heading('Documents Not Supplied', level=2)
		# p = self.document.add_paragraph()
		# run = p.add_run()
		p = self.document.add_paragraph('Aircraft Maintenace Manual (AMM)', style='Bullet Point Level 2')
		p = self.document.add_paragraph('Aircraft Illustrated Part Catalogue (IPC)', style='Bullet Point Level 2')

		self.document.add_heading('Documents Supplied', level=2)

		self.write_table(['Document No.', 'Issue', 'Title'], self.sb.drawings)
		
		self.document.add_heading('OEM Documents', level=2)
		
		self.write_table(['Title', 'Date', 'Ref', 'Version'], self.sb.oem_docs)

	
	def write_table(self, headers, data):
	
		table = self.document.add_table(rows=1, cols=len(headers)+1)
		
		#Add Headers
		hdr_cells = table.rows[0].cells
		for index, h in enumerate(headers):
			print(h)
			hdr_cells[index].text = h
		
			self.set_cell_border(hdr_cells[index], top=self.hdr_border,  bottom=self.hdr_border, start=self.hdr_border, end=self.hdr_border)
			
			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b1b5b2"/>'.format(nsdecls('w')))
			hdr_cells[index]._tc.get_or_add_tcPr().append(shading_elm_1)
			
		for row in data:
			new_row = table.add_row()
			row_cells = new_row.cells	
			
			for index, c in enumerate(row):
				run = row_cells[index].paragraphs[0].add_run(str(c))
				self.set_cell_border(row_cells[index], bottom = self.hdr_border, start=self.hdr_border, end=self.hdr_border)
				
				
	def set_cell_border(self, cell, **kwargs):
		"""
		Set cell`s border
		Usage:

		set_cell_border(
			cell,
			top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
			bottom={"sz": 12, "color": "#00FF00", "val": "single"},
			start={"sz": 24, "val": "dashed", "shadow": "true"},
			end={"sz": 12, "val": "dashed"},
		)
		"""
		tc = cell._tc
		tcPr = tc.get_or_add_tcPr()

		# check for tag existnace, if none found, then create one
		tcBorders = tcPr.first_child_found_in("w:tcBorders")
		if tcBorders is None:
			tcBorders = OxmlElement('w:tcBorders')
			tcPr.append(tcBorders)

		# list over all available tags
		for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
			edge_data = kwargs.get(edge)
			if edge_data:
				tag = 'w:{}'.format(edge)

				# check for tag existnace, if none found, then create one
				element = tcBorders.find(qn(tag))
				if element is None:
					element = OxmlElement(tag)
					tcBorders.append(element)

				# looks like order of attributes is important
				for key in ["sz", "val", "color", "space", "shadow"]:
					if key in edge_data:
						element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level