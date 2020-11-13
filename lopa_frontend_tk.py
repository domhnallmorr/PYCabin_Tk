import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter.ttk import *
import tkinter.messagebox

import LOPA_backend as lopa_bk
# from Pycabin_Backend import ipc_generator as ipc_gen
# from Pycabin_Backend import aircraft_models
import data_input_checks_tk

import gui_styles_tk
import components_tk
import double_scrollbar
import lopa_draw
import seats_draw
import docx_functions
#from Pycabin_Backend import lopa_draw_redo

import matplotlib
matplotlib.use('TkAgg')

from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from matplotlib.figure import Figure
import treeview_functions
import comment_box

#import ezdxf

from docx import Document
from docx.enum.text import WD_BREAK
#from ezdxf.tools.standards import linetypes  # some predefined line types

def check_lopa_used(self):
	lopa_used = False
	psus = []
	psu_dict = components_tk.get_all_components(self.mainapp, 'PSUs')
	
	for p in psu_dict['All']:
		
		if self.backend.title == self.mainapp.frames[p].backend.lopa:
			lopa_used = True
			psus.append(p)
			
	return lopa_used, psus

def get_seat_y_datum(station, aircraft, side):

	if aircraft in ['A320', 'A319']:
		
		y_datum = 30.12
	
	elif aircraft in ['B737-800']:
		y_datum = 24.755
		
	if side == 'LHS':
		y_datum = y_datum*-1
	
	return y_datum

class LOPA_Page_Tk(tk.Frame):

	def __init__(self, container, mainapp):
		tk.Frame.__init__(self, container)
		
		self.mainapp = mainapp

		self.top_label = tk.Label(self, text=('LOPA: '),font=self.mainapp.title_font, anchor="w")
		self.top_label.pack(fill=tk.BOTH, expand=True)
		
		self.backend = lopa_bk.LOPA_Backend(self, mainapp)
		
		self.treeview_iid = None
		self.setup_notebook()
		self.setup_scrollable_frames()
		self.setup_label_frames()
		self.setup_labels()
		self.setup_treeviews()
		self.setup_buttons()
		self.add_lopa_plot()
		self.add_overwing_plot()
		self.set_grid_configures()
		
	def setup_scrollable_frames(self):

		self.main_scroll_frame = double_scrollbar.Double_ScrollableFrame(self.main_tab, self.mainapp)
		self.main_scroll_frame.pack(fill=tk.BOTH, expand=True)
		
		self.lopa_scroll_frame = double_scrollbar.Double_ScrollableFrame(self.seat_tab, self.mainapp)
		self.lopa_scroll_frame.pack(fill=tk.BOTH, expand=True)
	
		self.weight_scroll_frame = double_scrollbar.Double_ScrollableFrame(self.weight_tab, self.mainapp)
		self.weight_scroll_frame.pack(fill=tk.BOTH, expand=True)
		
		self.overwing_scroll_frame = double_scrollbar.Double_ScrollableFrame(self.overwing_tab, self.mainapp)
		self.overwing_scroll_frame.pack(fill=tk.BOTH, expand=True)
		# '''
		# self.frame = scrollable_frame.ScrollableFrame(self.seat_tab)
		# self.frame.grid(row=0,column=0,stick='nsew')
		# self.lopa_frame = self.frame.scrollable_frame
		# self.seat_tab.grid_columnconfigure(0, weight=1)
		# self.seat_tab.grid_rowconfigure(0, weight=1)
		# #self.lopa_frame.grid_columnconfigure(4, weight=1)
		# '''
	def set_grid_configures(self):
	
		self.lopa_scroll_frame.inner.grid_columnconfigure(4, weight=1)
		self.overwing_scroll_frame.inner.grid_columnconfigure(0, weight=1)
		self.overwing_frame.grid_columnconfigure(7, weight=1)
		self.overwing_frame.grid_rowconfigure(3, weight=1)
		#self.weight_scroll_frame.frame.grid_columnconfigure(7, weight=1)
		pass
	def onFrameConfigure(self, canvas):
		'''Reset the scroll region to encompass the inner frame'''
		canvas.configure(scrollregion=canvas.bbox("all"))

	def FrameWidth(self, event):
		canvas_width = event.width
		self.lopa_frame_canvas.itemconfig(self.canvas_frame, width = canvas_width)
		
	def setup_notebook(self):
	
		self.note = ttk.Notebook(self)
		self.main_tab = Frame(self.note)
		self.seat_tab = Frame(self.note)
		self.weight_tab = Frame(self.note)
		self.overwing_tab = Frame(self.note)
		self.comments_tab = Frame(self.note)
		
		self.note.add(self.main_tab, text = "Main")
		self.note.add(self.seat_tab, text = "Seat Layout")
		self.note.add(self.weight_tab, text = "Seat Weights")
		self.note.add(self.overwing_tab, text = "Over Wing Exits")
		self.note.add(self.comments_tab, text = "Comments")
		
		#self.note.grid(row=1,column=0,sticky='NSEW')
		self.note.pack(fill=tk.BOTH, expand=True)
		# ####### COMMENTS TEXT ######################################
		self.comment_text = tk.Text(self.comments_tab, width = 110, height = 50, state='disabled')
		self.comment_text.grid(row=1, column=0, columnspan = 8, sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

	def setup_label_frames(self):
	
		self.main_frame = LabelFrame(self.main_scroll_frame.inner,text="LOPA Details:")
		self.main_frame.grid(row=2, column=0, columnspan = 16, rowspan = 2,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
		self.lav_frame = LabelFrame(self.main_scroll_frame.inner,text="Lavs:")
		self.lav_frame.grid(row=4, column=0, columnspan = 8, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

		self.galley_frame = LabelFrame(self.main_scroll_frame.inner,text="Galleys:")
		self.galley_frame.grid(row=5, column=0, columnspan = 8, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

		self.wb_frame = LabelFrame(self.main_scroll_frame.inner,text="Windbreakers:")
		self.wb_frame.grid(row=6, column=0, columnspan = 8, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
		self.items_frame = LabelFrame(self.main_scroll_frame.inner,text="Seat Item Numbers:")
		self.items_frame.grid(row=4, column=8, columnspan = 8, rowspan = 3,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
		self.seats_frame = LabelFrame(self.lopa_scroll_frame.inner,text="Seats:")
		self.seats_frame.grid(row=3, column=0, columnspan = 4, rowspan = 2,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
		self.preview_frame = LabelFrame(self.lopa_scroll_frame.inner,text="LOPA Preview:")
		self.preview_frame.grid(row=5, column=0, columnspan = 5, rowspan = 2,sticky='NSEW',padx=5, pady=5, ipadx=2, ipady=5)
		self.preview_frame.grid_columnconfigure(4, weight=1)
		
		weight_frame = self.weight_scroll_frame.inner
		self.weight_frame = LabelFrame(weight_frame,text="Seat Weights:")
		self.weight_frame.grid(row=3, column=0, columnspan = 4, rowspan = 2,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

		self.overwing_frame = LabelFrame(self.overwing_scroll_frame.inner, text='Overwing Exit Plot')
		self.overwing_frame.pack(fill=tk.BOTH, expand=True)
		#self.overwing_frame.grid(row =2, column=0, sticky='NW')
		
	def setup_labels(self):
			
		self.dwg_no_label = gui_styles_tk.create_label(self.main_frame,'')
		self.dwg_no_label.grid(row = 2, column = 0,columnspan=2, pady=2,padx=2, sticky="nsew")

		self.dwg_rev_label = gui_styles_tk.create_label(self.main_frame,'')
		self.dwg_rev_label.grid(row = 2, column = 2,columnspan=2, pady=2,padx=2, sticky="nsew")
		
		self.description_label = gui_styles_tk.create_label(self.main_frame,'')
		self.description_label.grid(row = 2, column = 4,columnspan=8, pady=2,padx=2, sticky="nsew")
		self.description_label.configure(width=70)
		
		self.no_lhs_label = gui_styles_tk.create_label(self.main_frame,'')
		self.no_lhs_label.grid(row = 3, column = 0,columnspan=2, pady=2,padx=2, sticky="nsew")

		self.no_rhs_label = gui_styles_tk.create_label(self.main_frame,'')
		self.no_rhs_label.grid(row = 3, column = 2,columnspan=2, pady=2,padx=2, sticky="nsew")

		self.ac_label = gui_styles_tk.create_label(self.main_frame,'')
		self.ac_label.grid(row = 3, column = 4,columnspan=2, pady=2,padx=2, sticky="nsew")
		
		
	def setup_treeviews(self):
	
		# self.monument_tree = ttk.Treeview(self.lav_frame, selectmode="extended",columns=("A","B",'C'),height = 10)
		# #self.monument_tree.grid(row=1,column=0, columnspan= 6,sticky="nsew")
		# self.monument_tree.heading("#0", text="#")
		# self.monument_tree.column("#0",minwidth=0,width=60, stretch='NO')
		# self.monument_tree.heading("A", text="Monument")	  
		# self.monument_tree.column("A",minwidth=0,width=200, stretch='NO') 
		# self.monument_tree.heading("B", text="Type")	  
		# self.monument_tree.column("B",minwidth=0,width=150, stretch='NO')
		# self.monument_tree.heading("C", text="Station (in)")	  
		# self.monument_tree.column("C",minwidth=0,width=130, stretch='NO')	

		# self.monument_tree.grid(row = 2, column = 0, columnspan = 8, sticky = 'NSEW')

		self.lav_tree = ttk.Treeview(self.lav_frame, selectmode="extended",columns=("A","B","C", "D"),height = 3)
		self.lav_tree.heading("#0", text="Lav")
		self.lav_tree.column("#0",minwidth=0,width=80, stretch='NO')
		self.lav_tree.heading("A", text="Installed")	  
		self.lav_tree.column("A",minwidth=0,width=100, stretch='NO') 
		self.lav_tree.heading("B", text="Station (in)")	  
		self.lav_tree.column("B",minwidth=0,width=130, stretch='NO')
		self.lav_tree.heading("C", text="CAS Installed")	  
		self.lav_tree.column("C",minwidth=0,width=130, stretch='NO')		
		self.lav_tree.heading("D", text="Doghouse Installed")	  
		self.lav_tree.column("D",minwidth=0,width=150, stretch='NO')
		self.lav_tree.grid(row = 2, column = 0, columnspan = 8, sticky = 'NSEW')
		self.lav_tree.bind("<Double-1>", lambda event, type='Lav', mode='edit': self.add_monument(event, type, mode))
		
		self.galley_tree = ttk.Treeview(self.galley_frame, selectmode="extended",columns=("A","B"),height = 2)
		self.galley_tree.heading("#0", text="Galley")
		self.galley_tree.column("#0",minwidth=0,width=100, stretch='NO')
		self.galley_tree.heading("A", text="Installed")	  
		self.galley_tree.column("A",minwidth=0,width=200, stretch='NO') 
		self.galley_tree.heading("B", text="Station (in)")	  
		self.galley_tree.column("B",minwidth=0,width=150, stretch='NO')
		self.galley_tree.grid(row = 2, column = 0, columnspan = 8, sticky = 'NSEW')

		self.wb_tree = ttk.Treeview(self.wb_frame, selectmode="extended",columns=("A"),height = 2)
		self.wb_tree.heading("#0", text="Windbreaker")
		self.wb_tree.column("#0",minwidth=0,width=200, stretch='NO')
		self.wb_tree.heading("A", text="Station (in)")	  
		self.wb_tree.column("A",minwidth=0,width=150, stretch='NO') 
		self.wb_tree.grid(row = 2, column = 0, columnspan = 8, sticky = 'NSEW')
		self.wb_tree.bind("<Double-1>", lambda event, type='Windbreaker', mode='edit': self.add_monument(event, type, mode))
		
		self.item_tree = ttk.Treeview(self.items_frame, selectmode="extended",columns=("A","B",'C'),height = 10)
		#self.monument_tree.grid(row=1,column=0, columnspan= 6,sticky="nsew")
		self.item_tree.heading("#0", text="#")
		self.item_tree.column("#0",minwidth=0,width=60, stretch='NO')
		self.item_tree.heading("A", text="Seat")	  
		self.item_tree.column("A",minwidth=0,width=200, stretch='NO') 
		self.item_tree.heading("B", text="Item Number")	  
		self.item_tree.column("B",minwidth=0,width=150, stretch='NO')
		self.item_tree.heading("C", text="Qty")	  
		self.item_tree.column("C",minwidth=0,width=130, stretch='NO')	

		self.item_tree.grid(row = 2, column = 0, columnspan = 8, sticky = 'NSEW')
		
		tree_height = 8
		self.LHS_lopa_tree = ttk.Treeview(self.seats_frame,selectmode="extended",columns=("A","B",'C'),height = tree_height)
		self.LHS_lopa_tree.grid(row=1,column=0, sticky="nsew")
		self.LHS_lopa_tree.heading("#0", text="Row")
		self.LHS_lopa_tree.column("#0",minwidth=0,width=50, stretch='NO')
		self.LHS_lopa_tree.heading("A", text="Seat")   
		self.LHS_lopa_tree.column("A",minwidth=0,width=250, stretch='NO') 
		self.LHS_lopa_tree.heading("B", text="Pitch")   
		self.LHS_lopa_tree.column("B",minwidth=0,width=60, stretch='NO')
		self.LHS_lopa_tree.heading("C", text="Station")   
		self.LHS_lopa_tree.column("C",minwidth=0,width=150, stretch='NO')

		LHS_lopa_tree_scrollbar = Scrollbar(self.seats_frame, command=self.LHS_lopa_tree.yview)
		LHS_lopa_tree_scrollbar.grid(row=1, column=1, sticky='nsew')
		self.LHS_lopa_tree.config(yscrollcommand=LHS_lopa_tree_scrollbar.set)
		
		self.LHS_lopa_tree.bind("<Double-1>", lambda event, side='LHS': self.seat_double_click(event, side))
		
		self.RHS_lopa_tree = ttk.Treeview(self.seats_frame,selectmode="extended",columns=("A","B",'C'),height = tree_height)
		self.RHS_lopa_tree.grid(row=1,column=2, sticky="nsew")
		self.RHS_lopa_tree.heading("#0", text="Row")
		self.RHS_lopa_tree.column("#0",minwidth=0,width=50, stretch='NO')
		self.RHS_lopa_tree.heading("A", text="Seat")   
		self.RHS_lopa_tree.column("A",minwidth=0,width=250, stretch='NO') 
		self.RHS_lopa_tree.heading("B", text="Pitch")   
		self.RHS_lopa_tree.column("B",minwidth=0,width=60, stretch='NO')
		self.RHS_lopa_tree.heading("C", text="Station")   
		self.RHS_lopa_tree.column("C",minwidth=0,width=150, stretch='NO')

		RHS_lopa_tree_scrollbar = Scrollbar(self.seats_frame, command=self.RHS_lopa_tree.yview)
		RHS_lopa_tree_scrollbar.grid(row=1, column=3, sticky='nsew')
		self.RHS_lopa_tree.config(yscrollcommand=RHS_lopa_tree_scrollbar.set)

		self.RHS_lopa_tree.bind("<Double-1>", lambda event, side='RHS': self.seat_double_click(event, side))
		
		# ___ WEIGHTS ____
		self.LHS_weight_tree = ttk.Treeview(self.weight_frame,selectmode="extended",columns=("A","B",'C','D','E'),height=30)
		self.LHS_weight_tree.grid(row=0,rowspan = 16,column=0,columnspan = 2, sticky="nsew")
		self.LHS_weight_tree.heading("#0", text="Row")
		self.LHS_weight_tree.column("#0",minwidth=0,width=50, stretch='NO')
		self.LHS_weight_tree.heading("A", text="Seat")   
		self.LHS_weight_tree.column("A",minwidth=0,width=160, stretch='NO') 
		self.LHS_weight_tree.heading("B", text="Pitch")   
		self.LHS_weight_tree.column("B",minwidth=0,width=50)
		self.LHS_weight_tree.heading("C", text="Weight")   
		self.LHS_weight_tree.column("C",minwidth=0,width=50)
		self.LHS_weight_tree.heading("D", text="Allowable")   
		self.LHS_weight_tree.column("D",minwidth=0,width=80)
		self.LHS_weight_tree.heading("E", text="Delta")   
		self.LHS_weight_tree.column("E",minwidth=0,width=80)     

		LHS_weight_tree_scrollbar = Scrollbar(self.weight_frame, command=self.LHS_weight_tree.yview)
		LHS_weight_tree_scrollbar.grid(row=0, column=2, rowspan = 16,sticky='nsew')
		self.LHS_weight_tree.config(yscrollcommand=LHS_weight_tree_scrollbar.set)
		
		#setup RHS treeview
		self.RHS_weight_tree = ttk.Treeview(self.weight_frame,selectmode="extended",columns=("A","B",'C','D','E'),height=30)
		self.RHS_weight_tree.grid(row=0,rowspan = 16,column=3,columnspan = 3 , sticky="nsew")
		self.RHS_weight_tree.heading("#0", text="Row")
		self.RHS_weight_tree.column("#0",minwidth=0,width=50, stretch='NO')
		self.RHS_weight_tree.heading("A", text="Seat")   
		self.RHS_weight_tree.column("A",minwidth=0,width=160, stretch='NO') 
		self.RHS_weight_tree.heading("B", text="Pitch")   
		self.RHS_weight_tree.column("B",minwidth=0,width=50)
		self.RHS_weight_tree.heading("C", text="Weight")   
		self.RHS_weight_tree.column("C",minwidth=0,width=50)
		self.RHS_weight_tree.heading("D", text="Allowable")   
		self.RHS_weight_tree.column("D",minwidth=0,width=80)
		self.RHS_weight_tree.heading("E", text="Delta")   
		self.RHS_weight_tree.column("E",minwidth=0,width=80)    		

		RHS_weight_tree_scrollbar = Scrollbar(self.weight_frame, command=self.RHS_weight_tree.yview)
		RHS_weight_tree_scrollbar.grid(row=0, column=6, rowspan = 16,sticky='nsew')
		self.RHS_weight_tree.config(yscrollcommand=RHS_weight_tree_scrollbar.set)
		
		#add in entries for life vest and literature weights
		self.weight_description=Label(self.weight_frame,text="Weights below are per Passenger",width=25)
		self.weight_description.grid(row=0,column=7, sticky="nsew")
		
		self.life_vest_weight = tk.StringVar()
		self.life_vest_weight_label=Label(self.weight_frame,text="Life Vest Weight (lbs)",width=25)
		self.life_vest_weight_label.grid(row=2,column=7, sticky="nsew")
		self.life_vest_weight_entry=Entry(self.weight_frame, width=20, textvariable=self.life_vest_weight)
		self.life_vest_weight_entry.insert(END, '2')
		self.life_vest_weight_entry.grid(row=3,column=7, sticky="nsew")

		self.literature_weight = tk.StringVar()
		self.literature_weight_label=Label(self.weight_frame,text="Literature Weight (lbs)",width=25)
		self.literature_weight_label.grid(row=4,column=7, sticky="nsew")
		self.literature_weight_entry=Entry(self.weight_frame, width=20, textvariable=self.literature_weight)
		self.literature_weight_entry.insert(END, '3')
		self.literature_weight_entry.grid(row=5,column=7, sticky="nsew")		

		#configure tags to color treeview
		self.LHS_weight_tree.tag_configure('bad', background='orange')
		self.LHS_weight_tree.tag_configure('good', background='green')
		self.RHS_weight_tree.tag_configure('bad', background='orange')
		self.RHS_weight_tree.tag_configure('good', background='green')        
		

		# Update button
		self.u=Button(self.weight_frame,text='Update Table', command= lambda: self.update_weight_table())
		self.u.grid(row=6,column=7, sticky="nsew")
		
	def setup_buttons(self):

		self.edit_btn = Button(self.main_scroll_frame.inner, text = 'Edit', image = self.mainapp.edit_icon2, compound = LEFT, width = 30, command= lambda: self.edit())
		self.edit_btn.grid(row=1, column=0, columnspan = 1, sticky='W',padx=5, pady=2, ipadx=2, ipady=2)

		# self.dxf_btn = Button(self.main_tab, text = 'Export to DXF',width = 30, command= lambda: self.export_dxf())
		# self.dxf_btn.grid(row=1, column=1, columnspan = 1, sticky='W',padx=5, pady=2, ipadx=2, ipady=2)

		self.ms_word_btn = Button(self.main_scroll_frame.inner, text = 'Export to Word', image = self.mainapp.word_icon2, compound = LEFT, width = 30, command= lambda: self.export_word())
		self.ms_word_btn.grid(row=1, column=2, columnspan = 1, sticky='W',padx=5, pady=2, ipadx=2, ipady=2)
		
		self.dxf_btn = Button(self.main_scroll_frame.inner, text = 'Export to DXF', image = self.mainapp.cad_icon2, compound = LEFT, width = 30, command= lambda: self.export_dxf())
		self.dxf_btn.grid(row=1, column=3, columnspan = 1, sticky='W',padx=5, pady=2, ipadx=2, ipady=2)
		
		self.add_wb_btn = Button(self.wb_frame, text = 'Add', image = self.mainapp.add_icon2, compound = LEFT,
								command =  lambda event=None, type='Windbreaker', mode='new': self.add_monument(event, type, mode))
		self.add_wb_btn.grid(row = 1, column = 0, columnspan = 2, sticky = 'NSEW')

		self.del_wb_btn = Button(self.wb_frame, text = 'Delete', image = self.mainapp.del_icon2, compound = LEFT,
								command =  lambda type='Windbreaker': self.del_windbreaker())
		self.del_wb_btn.grid(row = 1, column = 2, columnspan = 2, sticky = 'NSEW')
		
		self.edit_item_btn = Button(self.items_frame, text = 'Edit',
								command = self.edit_seat_item_nos)
		self.edit_item_btn.grid(row = 1, column = 0, columnspan = 2, sticky = 'NSEW')
		
		self.expand_lopa_tree_btn = Button(self.lopa_scroll_frame.inner, text = "Expand Trees",
							  command = lambda height=30, trees = [self.LHS_lopa_tree,self.RHS_lopa_tree]: self.expand_tree(trees,height))
		self.expand_lopa_tree_btn.grid(row=0, column=0, columnspan=4, sticky='nsew')

		self.overwing_plot_btn = Button(self.overwing_frame, text = "Update Plot",
									command= self.update_overwing_plot)
		self.overwing_plot_btn.grid(row=2, column=0, sticky='nsew', padx=2)

		self.overwing_combo = ttk.Combobox(self.overwing_frame, values=['LHS', 'RHS'], state='readonly')
		self.overwing_combo.grid(row=2, column=1, sticky='nsew', padx=2)
		
		self.edit_comment_button=Button(self.comments_tab,text='Edit', image = self.mainapp.edit_icon2, compound = LEFT,
										command= lambda self=self :comment_box.edit_comments(self))
		self.edit_comment_button.grid(row=0,column=0, pady=5,sticky="nsew", ipadx=2, ipady=2)
		
	def add_monument(self, event, type, mode):
		self.w=Add_Monument_Window(self, self.mainapp, self.master, type, mode)
		self.master.wait_window(self.w.top)	
		
		if self.w.button == 'ok':

			#treeview_functions.write_data_to_treeview(self.monument_tree, 'append', [data])
			
			self.update_component(self.w, 'edit')

	def del_windbreaker(self):
		
		#create temp class for update_component method
		
		tmp_class = lopa_bk.LOPA_Saved_State(self.backend)
		
		selected_items = self.wb_tree.selection()        
		for selected_item in selected_items:          
			self.wb_tree.delete(selected_item)
	  
		tmp_class.windbreakers = treeview_functions.get_all_treeview_items(self.wb_tree)
		
		self.update_component(tmp_class, 'edit')
	def expand_tree(self,trees,height):
		#print(tree.configure(height))
		for tree in trees:
			if tree['height'] < 30:
				tree.configure(height = 30)
			else:
				tree.configure(height = 10)
		
	def update_label_text(self):
				
		self.top_label.config(text=f'LOPA: {self.backend.title}')
		self.description_label.config(text=f' Description: {self.backend.description}')
		self.dwg_no_label.config(text=f' Drawing Number: {self.backend.drawing_no}')
		self.dwg_rev_label.config(text=f' Drawing Revision: {self.backend.drawing_rev}')
		self.ac_label.config(text=f' Aircraft Type: {self.backend.aircraft_type}')
		self.no_lhs_label.config(text=f' Number of LHS Rows: {self.backend.no_lhs_seats}')
		self.no_rhs_label.config(text=f' Number of RHS Rows: {self.backend.no_rhs_seats}')

	def update_monuments_tree(self):
		
		#treeview_functions.write_data_to_treeview(self.monument_tree, 'replace', self.backend.monuments)
		treeview_functions.write_data_to_treeview(self.lav_tree, 'replace', self.backend.lavs)
		treeview_functions.write_data_to_treeview(self.galley_tree, 'replace', self.backend.galleys)
		treeview_functions.write_data_to_treeview(self.wb_tree, 'replace', self.backend.windbreakers)

	def update_component(self, window, type):
		
		if type != 'new':
			orig_title = self.backend.title
			lopa_used, psus = check_lopa_used(self)
		else:
			lopa_used = False
			
		self.backend.update_component(window, type)
		self.update_label_text()
		self.update_monuments_tree()

		if self.treeview_iid:
			self.mainapp.main_treeview.item(self.treeview_iid, text = self.backend.title)
			components_tk.component_renamed(self)
			
		treeview_functions.write_data_to_treeview(self.LHS_lopa_tree, 'replace', self.backend.seat_layout['LHS'])
		treeview_functions.write_data_to_treeview(self.RHS_lopa_tree, 'replace', self.backend.seat_layout['RHS'])
		
		self.update_lopa_plot()
		
		if lopa_used:
			for p in psus:
				p = self.mainapp.frames[p]
				p.backend.lopa = self.backend.title
				p.update_component(p.backend, 'lopa')
					
		
	def update_lopa_plot(self):

		self.backend.ax1.clear()
		self.backend.ax2.clear()
		self.backend.ax3.clear()
		#lopa_draw.draw_lopa(self, self.backend.ax2, 'matplotlib', [[],[0,0],[]], True)
		lopa_draw.draw_aircraft(self.backend, self.backend.ax2, 'matplotlib', [0,0])
		lopa_draw.draw_floor(self.backend, self.backend.ax1, 'matplotlib', [0,0])
		lopa_draw.draw_floor(self.backend, self.backend.ax3, 'matplotlib', [0,0])
		#lopa_draw_redo.draw_seat_tracks(self.backend, self.backend.ax2, 'matplotlib', [0,0])

		lopa_draw.draw_seats_side(self.backend, self.backend.ax3, 'matplotlib', [0,0], 'LHS')
		lopa_draw.draw_seats_side(self.backend, self.backend.ax1, 'matplotlib', [0,0], 'RHS')

		if self.backend.aircraft_type in ['A320', 'A319']:
			y_datum = 30.12
		elif self.backend.aircraft_type in ['B737-800']:
			y_datum = 24.755
		lopa_draw.draw_seats_top_down(self.backend, self.backend.ax2, 'matplotlib', [0,y_datum*-1], 'LHS')
		lopa_draw.draw_seats_top_down(self.backend, self.backend.ax2, 'matplotlib', [0,y_datum], 'RHS')
		lopa_draw.draw_windbreakers_top_down(self.backend, self.backend.ax2, 'matplotlib', [0,0])
		lopa_draw.draw_windbreakers_side(self.backend, [self.backend.ax1, self.backend.ax3], 'matplotlib', [0,0])
		
		lopa_draw.draw_lavs_top_down(self.backend, self.backend.ax2, 'matplotlib', [0,0])
		lopa_draw.draw_lavs_side(self.backend, [self.backend.ax1, self.backend.ax3], 'matplotlib', [0,0])
		
		lopa_draw.draw_galleys_top_down(self.backend, self.backend.ax2, 'matplotlib', [0,0])
		lopa_draw.draw_galleys_side(self.backend, [self.backend.ax1, self.backend.ax3], 'matplotlib', [0,0,0,0])
		self.canvas.draw()

	def add_overwing_plot(self):
	
		self.overwing_figure = Figure(figsize=(5,5), dpi=100)
		self.overwing_ax1 = self.overwing_figure.add_subplot(111, aspect='equal', adjustable='box') #RHS
		#self.overwing_ax2 = self.overwing_figure.add_subplot(212, aspect='equal', adjustable='box') #LHS

		self.overwing_canvas = FigureCanvasTkAgg(self.overwing_figure, self.overwing_frame)
		self.overwing_canvas.draw()
		#canvas.get_tk_widget().pack(side=tk.BOTTOM, fill=tk.BOTH, expand=True)
		self.overwing_canvas.get_tk_widget().grid(row = 3, column = 0, columnspan=8, pady=2,sticky="nsew")
		
	def add_lopa_plot(self):
		
		self.canvas = FigureCanvasTkAgg(self.backend.lopa_figure, self.preview_frame)
		self.canvas.draw()
		#canvas.get_tk_widget().pack(side=tk.BOTTOM, fill=tk.BOTH, expand=True)
		self.canvas.get_tk_widget().grid(row = 3, column = 0, columnspan=8, pady=2,sticky="nsew")	

		toolbarFrame = Frame(master=self.preview_frame)
		
		toolbar = NavigationToolbar2Tk(self.canvas, toolbarFrame)
		toolbar.update()	
		toolbarFrame.grid(row = 2, column = 0, columnspan=5, pady=2,sticky="nsew")
	
	def update_overwing_plot(self):
		side = self.overwing_combo.get()
		
		if side in ['LHS', 'RHS']:
			for canvas in [self.overwing_ax1]:
				canvas.clear()
				
				s = 668.15
				w = 24.3
				
				x = [s-(w/2), s-(w/2), s+(w/2), s+(w/2), s-(w/2)]
				y = [14.1, 14.1+41.16, 14.1+41.16, 14.1, 14.1]
				
				canvas.fill(x, y, color='grey', alpha=0.6)# linestyle='dashed')
				
				canvas.plot([s, s], [14.1+41.16, 14.1+41.16+5], color='black', linestyle='dashed')
				canvas.text(s, 14.1+41.16+6, 'STA 668.15', ha='center')
				
				s = 701.5
				x = [s-(w/2), s-(w/2), s+(w/2), s+(w/2), s-(w/2)]
				
				canvas.fill(x, y, color='grey', alpha=0.6)

				canvas.plot([s, s], [14.1+41.16, 14.1+41.16+5], color='black', linestyle='dashed')
				canvas.text(s, 14.1+41.16+6, 'STA 701.5', ha='center')
				
			#find nearest seats
			ow_rows = self.backend.find_overwing_seats() #gives indexs in backend.seat_layout of nearest seats to OW exits
			print(ow_rows)
			sides = {'LHS': self.overwing_ax1, 'RHS': self.overwing_ax1}
			
			#draw seats (and get stations to annotatate
			

			annotate_stations = {'Rear': [], 'Front': [], 'Top': []}
			for i, indx in enumerate(ow_rows[side]):
				seat = self.mainapp.frames[self.backend.seat_layout[side][indx][1]].backend
				canvas = sides[side]
				canvas_type = 'matplotlib'
				station = self.backend.seat_layout[side][indx][3]
				side_datum = [station, 0]
				seats_draw.economy_seat_generic_side_view(seat, canvas, canvas_type, side_datum)
				
				print(station)
				print(seat.length_fwd)
				annotate_stations['Front'].append(station - float(seat.length_fwd))
				annotate_stations['Rear'].append(station + float(seat.length_aft))
				annotate_stations['Top'].append(float(seat.height))
					
			#add annotations
			for i, r in enumerate(annotate_stations['Rear']):
				if r != annotate_stations['Rear'][-1]:
					f = annotate_stations['Front'][i+1]
					t = max([annotate_stations['Top'][i+1], annotate_stations['Top'][i]])
					
					if f-r >=3:
						canvas.annotate(s='', xy=(r,t+1), xytext=(f,t+1), arrowprops=dict(arrowstyle='<->'))
					else:
						canvas.annotate(s='', xy=(r,t+1), xytext=(r-5,t+1), arrowprops=dict(arrowstyle='->'))
						canvas.annotate(s='', xy=(f+5,t+1), xytext=(f,t+1), arrowprops=dict(arrowstyle='<-'))
					canvas.text(r + ((f-r)/2), t+2, f'{str(round(f-r, 1))}"', ha='center')
			
			#add floor
			x1,x2,y1,y2 = canvas.axis()
			canvas.plot([x1, x2],[0, 0], color='black')
			self.overwing_canvas.draw()
		
	def seat_double_click(self, event, side):
		if side == 'LHS':
			item = self.LHS_lopa_tree.selection()[0]
			tree = self.LHS_lopa_tree
		elif side == 'RHS':
			item = self.RHS_lopa_tree.selection()[0]
			tree = self.RHS_lopa_tree
			
		row_data = list(tree.item(item,"values"))
		row_data.insert(0, tree.item(item,"text"))
		
		index = tree.index(item)
		
		self.w=Double_Click_Seat_Window_Tk(self, self.mainapp, self.master, side, row_data)
		self.master.wait_window(self.w.top)
		
		if self.w.button == 'ok':
			
			self.w.seat_layout[side][index] = [self.w.row_no, self.w.seat, self.w.pitch, None]
			
			if self.w.row_no_changed:
				lopa_bk.LOPA_Backend.update_row_numbers(self.w, index, side, self.w.row_no)
			#lopa_bk.LOPA_Backend.recalculate_stations(self.w)

			#if side == 'LHS':
			self.backend.ax3.clear()
			#else:
			self.backend.ax1.clear()

			self.update_component(self.w, 'edit')

	def edit(self):

		self.w= Edit_LOPA_Window_Tk(self.mainapp, self.master, None, 'edit', self)
		self.master.wait_window(self.w.top)	
			
		if self.w.button == 'ok':
			self.update_component(self.w, 'edit')
			
	def export_dxf(self):
		lopa_draw.gen_dxf(self)
		
	def export_word(self):

		mode = 'edit'
		
		self.w=Export_Word_Window(self.mainapp, self.master, mode, self)
		self.master.wait_window(self.w.top)	
		# document = Document()
		
		# ipc_data = self.gen_ipc_data('FROMTO')
		
		# ipc_gen.gen_ipc_table(document, ipc_data)
		
		# document.save(r'C:\Users\domhn\Documents\Python\Pycabin_Tkinter\V0.08\ipc.docx')

	def gen_ipc_data(self, fromto):
		data = treeview_functions.get_all_treeview_items(self.item_tree)
		ipc_data = []
		for d in data:
			desc = self.mainapp.frames[d[1]].backend.description
			ipc_data.append([d[2], d[1], desc, fromto, d[3]])	
		
		return ipc_data

	def update_weight_table(self):
		data_ok = True
		try:
			lvw = float(self.life_vest_weight_entry.get())
			lw = float(self.literature_weight_entry.get())
		except:	
			data_ok = False
			
		if data_ok:
			weight_data = {'LHS': [], 'RHS': []}
			
			for side in ['LHS', 'RHS']:
				for row in self.backend.seat_layout[side]:
					seat = self.mainapp.frames[row[1]]
					weight = round(seat.backend.calc_seat_weight(lvw, lw),2)
					weight_data[side].append([row[0], row[1], row[2], weight, '-', '-'])
					
		treeview_functions.write_data_to_treeview(self.LHS_weight_tree, 'replace',weight_data['LHS']) 
		treeview_functions.write_data_to_treeview(self.RHS_weight_tree,'replace', weight_data['RHS']) 

	def autogen_items(self):
	
		seats = {'LHS': {}, 'RHS': {}}
		numbers = {'LHS': 5, 'RHS': 6}
		
		#get quantities
		for side in ['LHS', 'RHS']:
			for row in self.backend.seat_layout[side]:
				if row[1] not in seats[side].keys():
					
					seats[side][row[1]] = 1
					
				else:
					seats[side][row[1]] += 1

		#Assign item numbers
		item_numbers = {}
		for side in ['LHS', 'RHS']:
			for seat in seats[side]:
				item_numbers[numbers[side]] = [seat, numbers[side], seats[side][seat]]
				numbers[side] += 2
		
		tree_list = []
		
		for index, s in enumerate(sorted(list(item_numbers.keys()))):
			
			tree_list.append(item_numbers[s])
			tree_list[-1].insert(0, index+1)#
			
		treeview_functions.write_data_to_treeview(self.item_tree, 'replace', tree_list)	
	
	def edit_seat_item_nos(self):
		
		mode = 'edit'
		
		self.w=Edit_Item_Window(self.mainapp, self.master, mode, self)
		self.master.wait_window(self.w.top)	
		
		if self.w.button == 'ok':
			treeview_functions.write_data_to_treeview(self.item_tree, 'replace',self.w.data) 
		
class Edit_LOPA_Window_Tk(object):
	def __init__(self, mainapp, master, ac, mode, parent_lopa):
		#self.drawing_dictionary = drawing_dictionary
		top=self.top=Toplevel(master)
		top.grab_set()
		self.mainapp = mainapp
		self.mode = mode
		self.parent_lopa = parent_lopa

		self.orig_part_no = None
		
		if mode == 'edit':
			self.orig_part_no = parent_lopa.backend.title
			
		lopa_bk.setup_variables(self)
		if self.mode == 'edit':
			lopa_bk.update_variables(self, self.parent_lopa.backend)
		
		self.seats_dict = components_tk.get_all_components(mainapp, 'Seats')

		#if self.mode == 'edit':
			#self.part_no = parent_lopa.backend.part_no
		self.data_checks = {}
		self.setup_label_frames()
		self.setup_widgets()

		if self.mode == 'edit':
			self.set_default_values()
		
		if self.mode == 'new':
			self.title_entry.insert(0, 'A320 LOPA')
			#self.aircraft_combo.set('A320')
			self.lhs_pitch_combo.set(28)
			self.rhs_pitch_combo.set(28)
			self.lhs_rows_combo.set(30)
			self.rhs_rows_combo.set(30)
			self.lhs_seat_combo.set('A320 Seat 1')
			self.rhs_seat_combo.set('A320 Seat 2')
	def setup_label_frames(self):
		self.details_frame = LabelFrame(self.top,text="LOPA Details:")
		self.details_frame.grid(row=2, column=0, columnspan = 8, rowspan = 4,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

		self.seats_frame = LabelFrame(self.top,text="Seat Layout:")
		self.seats_frame.grid(row=6, column=0, columnspan = 8, rowspan = 2,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
	def setup_widgets(self):
		
		#if self.mode == 'edit':
		#	labels = ['Title:', 'Drawing Number:', 'Revision']
		#else:
		if self.mode == 'edit':
			state = 'disabled'
		else:
			state = 'normal'
		labels = ['Title:', 'Description:', 'Drawing Number:', 'Revision', 'Aircraft:', 'Row 13 Included:']
		row = 1
		gui_styles_tk.create_multiple_labels(self.details_frame, labels, row, 2, 20, 2, 2)

		self.title_entry=Entry(self.details_frame, width=20)		
		self.title_entry.grid(row=1,column=3,padx=2, pady=2,sticky = 'NSEW')
		self.data_checks['Title'] = ['title', self.title_entry, self.orig_part_no]
		
		self.description_entry=Entry(self.details_frame, width=50)		
		self.description_entry.grid(row=2,column=3,padx=2, pady=2,sticky = 'NSEW')
		
		self.drawing_entry=Entry(self.details_frame, width=20)		
		self.drawing_entry.grid(row=3,column=3,padx=2, pady=2,sticky = 'NSEW')

		self.revision_combo= ttk.Combobox(self.details_frame, values=[i for i in range(99)])
		self.revision_combo.grid(row=4,column=3,padx=2, pady=2,sticky = 'NSEW')
		#self.revision_combo.set(1)
			

		#if self.mode != 'edit':
			
		#self.aircraft_combo= ttk.Combobox(self.details_frame, values=['A320', 'A319', 'B737-800'],state='readonly')
		self.aircraft_combo= ttk.Combobox(self.details_frame, values=['A320'],state='readonly')
		self.aircraft_combo.grid(row=5,column=3,padx=2, pady=2,sticky = 'NSEW')
		self.aircraft_combo.bind("<<ComboboxSelected>>", self.aircraft_selected)
		if self.mode == 'edit':
			self.aircraft_combo.set(self.parent_lopa.backend.aircraft_type)
			self.aircraft_combo.config(state='disabled')
		else:
			self.data_checks['Aircraft Type'] = ['combo', self.aircraft_combo, 'in values', 'Aircraft Type']
		
		self.economy_entry=Entry(self.details_frame, width=20)		
		#self.economy_entry.grid(row=6,column=3,padx=2, pady=2,sticky = 'NSEW')		

		self.row13_ckbx = Checkbutton(self.details_frame, text="")
		self.row13_ckbx.grid(row=7,column=3,padx=2, pady=2,sticky = 'NSEW')
		
		labels = ['Number Rows LHS:', 'Default Seat LHS:', 'Default Pitch LHS (in):']
		row = 2
		gui_styles_tk.create_multiple_labels(self.seats_frame, labels, row, 2, 20, 2, 2)

		self.lhs_rows_combo= ttk.Combobox(self.seats_frame, values=[i for i in range(31)],state=state)
		self.lhs_rows_combo.grid(row=2,column=3,padx=2, pady=2,sticky = 'NSEW')
		if self.mode != 'edit':
			self.data_checks['Number of LHS Seats'] = ['combo', self.lhs_rows_combo,  'in values', 'LHS No. of Rows']
		
		self.lhs_seat_combo= ttk.Combobox(self.seats_frame,state=state)
		self.lhs_seat_combo.grid(row=3,column=3,padx=2, pady=2,sticky = 'NSEW')
		if self.mode != 'edit':
			self.data_checks['LHS Default Seat'] = ['combo', self.lhs_seat_combo, 'in values', 'LHS Seat']
		
		self.lhs_pitch_combo= ttk.Combobox(self.seats_frame, values=[28, 29, 30, 31, 32, 33],state=state)
		self.lhs_pitch_combo.grid(row=4,column=3,padx=2, pady=2,sticky = 'NSEW')
		if self.mode != 'edit':
			self.data_checks['LHS Default Pitch'] = ['combo', self.lhs_pitch_combo, 'int greater than 27', 'LHS Pitch']
		
		labels = ['Number Rows RHS:', 'Default Seat RHS:', 'Default Pitch RHS (in):']
		row = 2
		gui_styles_tk.create_multiple_labels(self.seats_frame, labels, row, 4, 20, 2, 2)

		self.rhs_rows_combo= ttk.Combobox(self.seats_frame, values=[i for i in range(31)],state=state)
		self.rhs_rows_combo.grid(row=2,column=5,padx=2, pady=2,sticky = 'NSEW')
		if self.mode != 'edit':
			self.data_checks['Number of RHS Seats'] = ['combo', self.rhs_rows_combo, 'in values', 'RHS No. of Rows']
		
		self.rhs_seat_combo= ttk.Combobox(self.seats_frame,state=state)
		self.rhs_seat_combo.grid(row=3,column=5,padx=2, pady=2,sticky = 'NSEW')
		if self.mode != 'edit':
			self.data_checks['RHS Default Seat'] = ['combo', self.rhs_seat_combo, 'in values', 'RHS Seat']
		
		self.rhs_pitch_combo= ttk.Combobox(self.seats_frame, values=[28, 29, 30, 31, 32, 33],state=state)
		self.rhs_pitch_combo.grid(row=4,column=5,padx=2, pady=2,sticky = 'NSEW')
		if self.mode != 'edit':
			self.data_checks['RHS Default Pitch'] = ['combo', self.rhs_pitch_combo, 'int greater than 27', 'RHS Pitch']
		
		# ok button
		self.ok_button=Button(self.top,text='OK', command= lambda button = 'ok': self.cleanup(button))
		self.ok_button.grid(row=8,column=3, pady=5,sticky="nsew")

		# cancel button
		self.b=Button(self.top,text='Cancel', command= lambda button = 'cancel': self.cleanup(button))
		self.b.grid(row=8,column=4, pady=5,sticky="nsew")

		self.button = 'cancel'
		
	def set_default_values(self):
		
		self.title_entry.insert(0, self.parent_lopa.backend.title)
		self.description_entry.insert(0, self.parent_lopa.backend.description)
		self.drawing_entry.insert(0, self.parent_lopa.backend.drawing_no)
		self.revision_combo.set(self.parent_lopa.backend.drawing_rev)
		self.aircraft_combo.set(self.parent_lopa.backend.aircraft_type)
		self.lhs_rows_combo.set(self.parent_lopa.backend.no_lhs_seats)
		self.rhs_rows_combo.set(self.parent_lopa.backend.no_rhs_seats)

	def aircraft_selected(self, event):
		
		self.lhs_seat_combo.set('')
		self.rhs_seat_combo.set('')
		
		if self.aircraft_combo.get() in ['A320', 'A319']:
			self.lhs_seat_combo['values'] = [self.mainapp.frames[s].backend.title for s in self.seats_dict['A320 Family LHS']]
			self.rhs_seat_combo['values'] = [self.mainapp.frames[s].backend.title for s in self.seats_dict['A320 Family RHS']]
			
			if self.aircraft_combo.get() == 'A320':
				self.lhs_rows_combo['values'] = [i for i in range(31)]
				self.rhs_rows_combo['values'] = [i for i in range(31)]
			elif self.aircraft_combo.get()== 'A319':
				self.lhs_rows_combo['values'] = [i for i in range(26)]
				self.rhs_rows_combo['values'] = [i for i in range(26)]
			
		elif self.aircraft_combo.get() in ['B737-800']:
			self.lhs_seat_combo['values'] = [self.mainapp.frames[s].backend.title for s in self.seats_dict['B737 Family LHS']]
			self.rhs_seat_combo['values'] = [self.mainapp.frames[s].backend.title for s in self.seats_dict['B737 Family RHS']]	
			
			if self.aircraft_combo.get() == 'B737-800':
				self.lhs_rows_combo['values'] = [i for i in range(33)]
				self.rhs_rows_combo['values'] = [i for i in range(33)]
				
	def cleanup(self,button):
	
		if button == 'ok':
			#checks
			data_good, msg = data_input_checks_tk.check_data_input(self.data_checks, self.mainapp)
			
			if data_good:
				self.button = 'ok'
				
				if self.mode != 'edit':
					self.seat_layout = {'LHS': [], 'RHS': []}
					
					station = 0
					
					for i in range(int(self.lhs_rows_combo.get())):
						if i == 0:
							pitch = 340
						else:
							pitch = int(self.lhs_pitch_combo.get())
						station += pitch
						self.seat_layout['LHS'].append([i+1, self.lhs_seat_combo.get(), pitch, station])

					station = 0
					
					for i in range(int(self.rhs_rows_combo.get())):
						if i == 0:
							pitch = 340
						else:
							pitch = int(self.rhs_pitch_combo.get())
						station += pitch
						self.seat_layout['RHS'].append([i+1, self.rhs_seat_combo.get(), pitch, station])
				
				
				self.title = self.title_entry.get()
				self.description = self.description_entry.get()
				self.aircraft_type = self.aircraft_combo.get()
				self.no_lhs_seats = self.lhs_rows_combo.get()
				self.no_rhs_seats = self.rhs_rows_combo.get()				
				self.drawing_no = self.drawing_entry.get()
				self.drawing_rev = self.revision_combo.get()
				
				if self.mode == 'new':
					self.lavs, self.galleys = lopa_bk.set_default_monumnets(self.aircraft_type)
				
				self.top.destroy()
			
			else:
				tkinter.messagebox.showerror(master=self.top, title='Error', message=msg)
		else:
		
			self.top.destroy()
			
class Double_Click_Seat_Window_Tk(object):			
	def __init__(self, lopa, mainapp, master, side, row_data):
		top=self.top=Toplevel(master)
		top.grab_set()
		
		self.mainapp = mainapp
		self.side = side
		self.lopa = lopa
		self.row_data = row_data
		
		lopa_bk.setup_variables(self)
		lopa_bk.update_variables(self, self.lopa.backend)
		
		#self.seats, self.seats_dict = components_tk.get_all_seats(mainapp)
		self.seats_dict = components_tk.get_all_components(mainapp, 'Seats')
		self.setup_label_frames()
		self.setup_widgets()
	
	def setup_label_frames(self):
	
		self.options_frame = LabelFrame(self.top,text="Options:")
		self.options_frame.grid(row=2, column=0, columnspan = 4, rowspan = 4,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

	def setup_widgets(self):
	
		labels = ['Row Number:', 'Seat Part Number:', 'Pitch (in):',]
		row = 2
		gui_styles_tk.create_multiple_labels(self.options_frame, labels, row, 2, 20, 2, 2)
		
		if self.side == 'LHS':
			if self.lopa.backend.aircraft_type in ['A320', 'A319']:
				key = 'A320 Family LHS'
			elif self.lopa.backend.aircraft_type in ['B737-800']:
				key = 'B737 Family LHS'
				
		elif self.side == 'RHS':
			if self.lopa.backend.aircraft_type in ['A320', 'A319']:
				key = 'A320 Family RHS'
			elif self.lopa.backend.aircraft_type in ['B737-800']:
				key = 'B737 Family RHS'
				
		self.row_entry=Entry(self.options_frame, width=20)		
		self.row_entry.grid(row=2,column=3,padx=2, pady=2,sticky = 'NSEW')	
		self.row_entry.insert(0,self.row_data[0])
		
		self.seat_combo= ttk.Combobox(self.options_frame, values=self.seats_dict[key])
		self.seat_combo.grid(row=3,column=3,padx=2, pady=2,sticky = 'NSEW')
		self.seat_combo.set(self.row_data[1])
		
		pitches = [28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40]
		self.pitch_combo= ttk.Combobox(self.options_frame, values=pitches)
		self.pitch_combo.grid(row=4,column=3,padx=2, pady=2,sticky = 'NSEW')
		self.pitch_combo.insert(0,self.row_data[2])
		# ok button
		self.ok_button=Button(self.top,text='OK', command= lambda button = 'ok': self.cleanup(button))
		self.ok_button.grid(row=8,column=1, pady=5,sticky="nsew")

		# cancel button
		self.b=Button(self.top,text='Cancel', command= lambda button = 'cancel': self.cleanup(button))
		self.b.grid(row=8,column=2, pady=5,sticky="nsew")

		self.button = 'cancel'	

	def cleanup(self,button):
	
		if button == 'ok':
			self.button = 'ok'
			
			self.row_no = int(self.row_entry.get())
			self.seat = self.seat_combo.get()
			self.pitch = self.pitch_combo.get()
			self.top.destroy()
			
			if self.row_no != self.row_data[0]:
				self.row_no_changed = True
			else:
				self.row_no_changed = False
		else:
			self.top.destroy()
			
class Add_Monument_Window():
	def __init__(self, lopa, mainapp, master,monument_type, mode):
	
		top=self.top=Toplevel(master)
		top.grab_set()
		
		self.mainapp = mainapp

		self.lopa = lopa
		self.monument_type = monument_type
		self.mode = mode
		
		self.data_checks = {}
		
		lopa_bk.setup_variables(self)
		lopa_bk.update_variables(self, self.lopa.backend)
		
		if self.monument_type == 'Windbreaker':

			self.monuments = components_tk.get_all_components(mainapp, 'Windbreakers')
		
		elif self.monument_type == 'Lav' or self.monument_type == 'Galley':
			self.monuments = {'All': []}

		# self.lavs =components_tk.get_all_lavs(mainapp)
		# self.lavs = components_tk.gen_lav_dict(mainapp, self.lavs)
		
		self.setup_label_frames()
		self.setup_widgets()
		
		if mode == 'edit':
			if self.monument_type == 'Windbreaker':
				self.selected_item = self.lopa.wb_tree.selection()[0]     
				          
				current_part = self.lopa.wb_tree.item(self.selected_item, 'text')
				station = self.lopa.wb_tree.item(self.selected_item, 'values')[0]
				self.monument_combo.set(current_part)
				self.station_entry.insert(0, station)
				
				self.tree_index = 0
				for item in self.lopa.wb_tree.get_children():
					if item == self.selected_item:
						break
					self.tree_index += 1
			
			elif self.monument_type == 'Lav':
				self.selected_item = self.lopa.lav_tree.selection()[0] 
				
				current_part = self.lopa.lav_tree.item(self.selected_item, 'text')
				installed = self.lopa.lav_tree.item(self.selected_item, 'values')[0]
				station = self.lopa.lav_tree.item(self.selected_item, 'values')[1]
				cas_installed = self.lopa.lav_tree.item(self.selected_item, 'values')[2]
				doghouse_installed = self.lopa.lav_tree.item(self.selected_item, 'values')[3]
				
				self.monument_combo.set(current_part)
				self.monument_combo.config(state='disabled')
				self.installed_combo.set(installed)
				self.station_entry.insert(0, station)
				self.station_entry.config(state='readonly')
				self.cas_combo.set(cas_installed)
				self.doghouse_combo.set(doghouse_installed)
				
				if current_part == 'Lav A':
					self.doghouse_combo.config(state='disabled')
				
				
				self.tree_index = 0
				for item in self.lopa.lav_tree.get_children():
					if item == self.selected_item:
						break
					self.tree_index += 1
					
	def setup_label_frames(self):
	
		self.options_frame = LabelFrame(self.top,text="Options:")
		self.options_frame.grid(row=2, column=0, columnspan = 4, rowspan = 4,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)	
		
	def setup_widgets(self):
	
		if self.monument_type == 'Windbreaker':
			labels = [f'{self.monument_type}:', 'Station (in):',]
		elif self.monument_type == 'Lav':
			labels = [f'{self.monument_type}:', 'Installed', 'Station (in):', 'CAS Installed', 'Doghouse Installed']
			
		row = 2
		gui_styles_tk.create_multiple_labels(self.options_frame, labels, row, 2, 20, 2, 2)
		
		row = 2
		# self.type_combo= ttk.Combobox(self.options_frame, values=['Windbreakers', 'Lavs'])
		# self.type_combo.grid(row=2,column=3,padx=2, pady=2,sticky = 'NSEW')
		# self.type_combo.bind("<<ComboboxSelected>>", self.type_selected)
		
		self.monument_combo= ttk.Combobox(self.options_frame, values=self.monuments['All'], state='readonly')
		self.monument_combo.grid(row=row,column=3,padx=2, pady=2,sticky = 'NSEW')

		row += 1
		
		if self.monument_type == 'Lav':
			self.installed_combo = ttk.Combobox(self.options_frame, values=['Yes', 'No'], state='readonly')
			self.installed_combo.grid(row=row,column=3,padx=2, pady=2,sticky = 'NSEW')
			row += 1
			
		self.station_entry=Entry(self.options_frame, width=20)		
		self.station_entry.grid(row=row,column=3,padx=2, pady=2,sticky = 'NSEW')
		self.data_checks['Station'] = ['entry', self.station_entry, 'float positive', 'Station']
		row += 1
		
		if self.monument_type == 'Lav':
			self.cas_combo = ttk.Combobox(self.options_frame, values=['Yes', 'No'], state='readonly')
			self.cas_combo.grid(row=row,column=3,padx=2, pady=2,sticky = 'NSEW')
			row += 1

			self.doghouse_combo = ttk.Combobox(self.options_frame, values=['Yes', 'No'], state='readonly')
			self.doghouse_combo.grid(row=row,column=3,padx=2, pady=2,sticky = 'NSEW')
			row += 1
			
			
		# ok button
		self.ok_button=Button(self.top,text='OK', command= lambda button = 'ok': self.cleanup(button))
		self.ok_button.grid(row=8,column=1, pady=5,sticky="nsew")

		# cancel button
		self.b=Button(self.top,text='Cancel', command= lambda button = 'cancel': self.cleanup(button))
		self.b.grid(row=8,column=2, pady=5,sticky="nsew")

		self.button = 'cancel'

	# def type_selected(self, event):
		# type = self.type_combo.get()
		# if type == 'Windbreakers':
			# if self.lopa.backend.aircraft_type in ['A320', 'A319']:
				# self.monument_combo['values'] = self.wbs['A320 LHS'] + self.wbs['A320 RHS']
				# self.monument_combo.set('')

		# if type == 'Lavs':
			# if self.lopa.backend.aircraft_type in ['A320', 'A319']:
				# self.monument_combo['values'] = self.lavs['A320 LHS'] + self.lavs['A320 RHS']
				# self.monument_combo.set('')
				
	def cleanup(self,button):
	
		self.button = button
		
		if self.button == 'ok':
			#self.check_data()
			data_good, msg = data_input_checks_tk.check_data_input(self.data_checks, self.mainapp)
			
			if data_good:
				monument =self.monument_combo.get()
				station =self.station_entry.get()
				

				if self.monument_type == 'Windbreaker':
					data = [monument, station]
					
					if self.mode == 'new':
						self.windbreakers.append(data)
					elif self.mode == 'edit':
						
						self.windbreakers[self.tree_index] = data
				
				if self.monument_type == 'Lav':
					installed = self.installed_combo.get()
					cas_installed = self.cas_combo.get()
					doghouse_installed = self.doghouse_combo.get()
					
					data = [monument, installed, station, cas_installed, doghouse_installed]
					
					if self.mode == 'edit':
						self.lavs[self.tree_index] = data
				self.top.destroy()
				
			else:
				tkinter.messagebox.showerror(master=self.top, title='Error', message=msg)
			
		else:
			self.top.destroy()
			
class Edit_Item_Window(object):
	def __init__(self, mainapp, master, mode, parent_lopa):
		#self.drawing_dictionary = drawing_dictionary
		top=self.top=Toplevel(master)
		top.grab_set()
		self.mainapp = mainapp
		self.mode = mode
		self.parent_lopa = parent_lopa

		self.setup_label_frames()
		
		self.seats = {'LHS': {}, 'RHS': {}}
		
		current_items = treeview_functions.get_all_treeview_items(self.parent_lopa.item_tree)
		for side in self.seats:
			for row in self.parent_lopa.backend.seat_layout[side]:
				if row[1] not in self.seats[side]:
					self.seats[side][row[1]] = Entry(self.items_frame)
					
					#insert current number in tree
					for c in current_items:
						if c[1] == row[1]:
							self.seats[side][row[1]].insert(0, c[2])

		self.setup_widgets()		
		
		self.button = 'cancel'
	def setup_label_frames(self):
	
		self.auto_frame = LabelFrame(self.top,text="Autogen Options:")
		self.auto_frame.grid(row=2, column=0, columnspan = 4, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)

		self.items_frame = LabelFrame(self.top,text="Item Numbers:")
		self.items_frame.grid(row=3,column=0, columnspan = 4, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)	
		

	def setup_widgets(self):
	
		# ____ AUTOGEN OPTIONS ____
		
		Button(self.auto_frame, text = 'Autogen Numbers', command=self.autogen).grid(row=0, column = 1, sticky='NSEW')
		# starting number
		Label(self.auto_frame, text = 'Starting Number:').grid(row=1, column = 1, sticky='NSEW')
		self.start_entry = Entry(self.auto_frame)
		self.start_entry.grid(row=1, column = 2, padx=1, pady=1, sticky='NSEW')
		
		# format
		Label(self.auto_frame, text = 'Format:').grid(row=2, column = 1, sticky='NSEW')
		self.format_combo = ttk.Combobox(self.auto_frame, values=['1, 2, 3, etc', '01, 02, 03, etc'], state='readonly')
		self.format_combo.grid(row=2, column = 2, padx=1, pady=1, sticky='NSEW')
		
		# ____ SEATS ____
		
		column = 1
		for side in self.seats:
			row = 2
			for seat in self.seats[side]:
				Label(self.items_frame, text = seat).grid(row=row, column=column, sticky='NSEW')
				self.seats[side][seat].grid(row=row, column=column+1, sticky='NSEW', padx = 2)
				row +=1
			column += 2

		# ok button
		self.ok_button=Button(self.top,text='OK', command= lambda button = 'ok': self.cleanup(button))
		self.ok_button.grid(row=8,column=1, pady=5,sticky="nsew")

		# cancel button
		self.b=Button(self.top,text='Cancel', command= lambda button = 'cancel': self.cleanup(button))
		self.b.grid(row=8,column=2, pady=5,sticky="nsew")			
			
	def cleanup(self, button):
	
		self.button = button
		
		if self.button == 'ok':
			
			count = 1
			self.data = []
			for side in self.seats:
				for seat in self.seats[side]:
					#get qty of installed seats
					qty = 0
					for row in self.parent_lopa.backend.seat_layout[side]:
						if row[1] == seat:
							qty += 1
					self.data.append([count, seat, self.seats[side][seat].get(), qty])
					count += 1
					
			self.top.destroy()
			
		else:
			
			self.top.destroy()
	
	def autogen(self):
	
		start_no = 1
		
		try:
			start_no = int(self.start_entry.get())
		except:
			pass
			
		lhs_no = start_no
		for seat in self.seats['LHS']:
			self.seats['LHS'][seat].delete(0, 'end')
			if self.format_combo.get() == '01, 02, 03, etc'and lhs_no < 10:
				text = f'0{lhs_no}'
			else:
				text = lhs_no
				
			self.seats['LHS'][seat].insert(0, text)
			
			lhs_no += 2
			
		rhs_no = start_no+1
		for seat in self.seats['RHS']:
			self.seats['RHS'][seat].delete(0, 'end')
			if self.format_combo.get() == '01, 02, 03, etc'and rhs_no < 10:
				text = f'0{rhs_no}'
			else:
				text = rhs_no
				
			self.seats['RHS'][seat].insert(0, text)
			
			rhs_no += 2

class Export_Word_Window(object):
	def __init__(self, mainapp, master, mode, parent_lopa):
		#self.drawing_dictionary = drawing_dictionary
		top=self.top=Toplevel(master)
		top.grab_set()
		self.mainapp = mainapp
		self.mode = mode
		self.parent_lopa = parent_lopa

		self.setup_label_frames()		
		self.setup_widgets()
		
		self.button = 'cancel'
	def setup_label_frames(self):
	
		self.main_frame = LabelFrame(self.top, text='Options:')
		self.main_frame.grid(row=2, column=0, columnspan = 4, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
	def setup_widgets(self):
	
		Label(self.main_frame, text='Output File:').grid(row=1, column = 1, ipadx=1, ipady=1, padx=2, pady=2, sticky='NW')
		
		self.file_entry = Entry(self.main_frame, width = 80)
		self.file_entry.grid(row=1, column = 2, ipadx=1, ipady=1, padx=2, pady=2, sticky='NW')
		
		self.file_entry.insert(0, r'C:/Users/domhn/Documents/Python/Pycabin_Tkinter/V0.08/test.docx')
		Button(self.main_frame, text='Browse', command=self.browse).grid(row=1, column = 3, ipadx=1, ipady=1, padx=2, pady=2, sticky='NW')

		# ok button
		self.ok_button=Button(self.top,text='OK', command= lambda button = 'ok': self.cleanup(button))
		self.ok_button.grid(row=8,column=1, pady=5,sticky="nsew")

		# cancel button
		self.b=Button(self.top,text='Cancel', command= lambda button = 'cancel': self.cleanup(button))
		self.b.grid(row=8,column=2, pady=5,sticky="nsew")	
		
	def browse(self):
		
		filename = filedialog.asksaveasfilename()
		
		self.file_entry.delete(0, 'end')
		self.file_entry.insert(0, filename)
		
	def cleanup(self, button):
	
		self.button = button
		
		if self.button == 'ok':
			
			self.filename = self.file_entry.get()
			#Checks
			msg = None
			# is extension .docx
			if self.filename[-5:] != '.docx':
				msg = 'Extension must be .docx'
				
			# 
			if msg:	
				tkinter.messagebox.showerror(master=self.top, title='Error', message=msg)
			else:
				self.document = Document()
				parts = self.parent_lopa.backend.gen_parts_table()
				#SB table
				docx_functions.write_table(self.document, ['Qty', 'Part Number', 'Description'],parts)
				
				p = self.document.add_paragraph()
				run = p.add_run()
				run.add_break(WD_BREAK.PAGE)
		
				# IPC Table
				ipc_table = []
				item_nos = treeview_functions.get_all_treeview_items(self.parent_lopa.item_tree)
				for p in parts:
					item_number = ''
					for i in item_nos:
						if p[1] == i[1]:
							item_number = i[2]
							
					ipc_table.append([item_number, p[1], p[2], '', p[0]])
				
				docx_functions.write_table(self.document, ['ITEM NO.', 'PART NUMBER', 'DESCRIPTION', 'FROMTO', 'QTY'],ipc_table)
				
				self.document.save(self.filename)
				self.top.destroy()
			
		else:
			
			self.top.destroy()