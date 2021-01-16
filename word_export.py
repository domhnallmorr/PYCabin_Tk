import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter.ttk import *

import tkinter.messagebox

import docx_functions

from docx import Document
from docx.enum.text import WD_BREAK

import openpyxl
import excel_functions
import file_menu
import pandas as pd

class Export_Word_Excel_Window(object):
	def __init__(self, mainapp, master, mode, parent_component, file_type='word'):
		#self.drawing_dictionary = drawing_dictionary
		top=self.top=Toplevel(master)
		top.grab_set()
		self.mainapp = mainapp
		self.mode = mode
		self.parent_component = parent_component
		self.file_type = file_type

		self.setup_label_frames()		
		self.setup_widgets()
		
		self.button = 'cancel'

		if 'json' not in file_type:
			self.component_type = parent_component.backend.type
		else:
			self.component_type = None


	def setup_label_frames(self):
	
		self.main_frame = LabelFrame(self.top, text='Options:')
		self.main_frame.grid(row=2, column=0, columnspan = 4, rowspan = 1,sticky='NW',padx=5, pady=5, ipadx=2, ipady=5)
		
	def setup_widgets(self):
	
		Label(self.main_frame, text='Output File:').grid(row=1, column = 1, ipadx=1, ipady=1, padx=2, pady=2, sticky='NW')
		
		self.file_entry = Entry(self.main_frame, width = 80)
		self.file_entry.grid(row=1, column = 2, ipadx=1, ipady=1, padx=2, pady=2, sticky='NW')
		
		if self.file_type == 'word':
			self.file_entry.insert(0, r'C:/Users/domhn/Documents/Python/Pycabin_Tkinter/V0.18/test.docx')
		else:
			self.file_entry.insert(0, r'C:/Users/domhn/Documents/Python/Pycabin_Tkinter/V0.18/test.xlsx')

		Button(self.main_frame, text='Browse', command=self.browse).grid(row=1, column = 3, ipadx=1, ipady=1, padx=2, pady=2, sticky='NW')

		# ok button
		self.ok_button=Button(self.top,text='OK', command= lambda button = 'ok': self.cleanup(button))
		self.ok_button.grid(row=8,column=1, padx=5, pady=5,sticky="ne")

		# cancel button
		self.b=Button(self.top,text='Cancel', command= lambda button = 'cancel': self.cleanup(button))
		self.b.grid(row=8,column=2, padx=5, pady=5,sticky="nw")	
		
	def browse(self):
		
		filename = filedialog.asksaveasfilename()
		
		self.file_entry.delete(0, 'end')
		self.file_entry.insert(0, filename)
		
	def cleanup(self, button):
	
		self.button = button
		
		if self.button == 'ok':
			
			self.filename = self.file_entry.get()
			#Checks

			msg = None
			if self.file_type == 'word':
				# is extension .docx
				if self.filename[-5:] != '.docx':
					msg = 'Extension must be .docx'
				
			elif self.file_type == 'excel':
				if self.filename[-5:] != '.xlsx':
					msg = 'Extension must be .xlsx'

			elif self.file_type == 'dxf':
				if self.filename[-4:] != '.dxf':
					msg = 'Extension must be .dxf'

			elif self.file_type == 'load json' or self.file_type == 'save json':
				if self.filename[-5:] != '.json':
					msg = 'Extension must be .json'

			# 
			if msg:	
				tkinter.messagebox.showerror(master=self.top, title='Error', message=msg)
			else:

				if self.file_type == 'word':
					saved = self.export_word_doc()

				elif self.file_type == 'excel':

					saved = self.export_excel_file()

				elif self.file_type == 'dxf':
					self.parent_component.draw_dxf(self.filename)
					saved=True

				elif self.file_type == 'load json':
					saved = self.load_json()

				if saved:
					self.top.destroy()
			
		else:
			
			self.top.destroy()

	def export_word_doc(self):

		self.document = Document()
		parts = self.parent_component.backend.gen_parts_table()
		#SB table
		docx_functions.write_table(self.document, ['Qty', 'Part Number', 'Description'],parts)
		
		p = self.document.add_paragraph()
		run = p.add_run()
		run.add_break(WD_BREAK.PAGE)

		ipc_table = self.parent_component.backend.gen_ipc_table()
		
		docx_functions.write_table(self.document, ['ITEM NO.', 'PART NUMBER', 'DESCRIPTION', 'FROMTO', 'QTY'],ipc_table)
		
		try:
			self.document.save(self.filename)
			saved = True
		except PermissionError:

			tkinter.messagebox.showerror(master=self.top, title='Error', message='Could not save file, permission denied')
			saved = False		

		return saved

	def export_excel_file(self):

		excel_data = self.parent_component.backend.gen_excel_data()

		if excel_data['type'] == 'list':

			wb = openpyxl.Workbook()
		
			sheet_styles = excel_functions.setup_styles()
			count = 1

			for table in excel_data['data']:

				wb.create_sheet(index=count, title=(table))
				wb.active = count
				sheet = wb.active

				excel_functions.add_data_to_sheet(wb, sheet, excel_data['data'][table], 1, 1, sheet_styles['Normal'])

				count += 1
	
			try:
				wb.save(self.filename)
				saved = True
			except PermissionError:

				tkinter.messagebox.showerror(master=self.top, title='Error', message='Could not save file, permission denied')
				saved = False

		elif excel_data['type'] == 'df':
			
			with pd.ExcelWriter(self.filename) as writer:
				for table in excel_data['data']:
					
					excel_data['data'][table].to_excel(writer, sheet_name=table)
					saved=True

	def load_json(self):

		try:
			file_menu.load_file(self.filename)
			saved=True
		except:
			tkinter.messagebox.showerror(master=self.top, title='Error', message='Failed to Load File')
			saved = False			


